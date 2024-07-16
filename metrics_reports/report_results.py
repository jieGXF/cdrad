import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import os
from sklearn.metrics import roc_curve, auc, accuracy_score, f1_score, recall_score, precision_score
from sklearn.metrics import confusion_matrix
from sklearn.calibration import calibration_curve
import seaborn as sns
from time import time
from docx.shared import Inches
from collections import OrderedDict


class ResultReportor(object):
    def __init__(self, test_pred_data_frame, train_pred_data_frame, save_dir, document, subj_column_name, label_column_name, do_feature_selection,deterministic=True):
        if deterministic:
            np.random.seed(12345)
        self.save_dir = None
        self.best_classifier = None
        self.pred_data_frame = None
        self.save_dir = save_dir
        self.train_pred_data_frame = train_pred_data_frame
        self.test_pred_data_frame = test_pred_data_frame
        self.document = document

        self.subj_column_name =  subj_column_name
        self.label_column_name = label_column_name

        self.do_feature_selection = do_feature_selection

        self.cutoff = None

    def choose_cutoff(self, prods, groudtruth):
        """
        based on Youden index
        :param prods:
        :param groudtruth:
        :return:
        """
        fpr, tpr, threshod = roc_curve(groudtruth, prods)
        criterion = tpr + (1 - fpr)
        threshod_best = threshod[np.argmax(criterion)]
        return threshod_best


    def plot_roc_curves(self,pred_data_frame, data_tag='train'):
        """
        :param pred_data_frame: the result after classification
        :return:
        """
        color_list = ['r', 'g', 'b']
        gt = pred_data_frame[self.label_column_name].values
        method_names = pred_data_frame.drop([self.subj_column_name, self.label_column_name], axis=1).columns.values
        method_metrics = OrderedDict()
        plt.figure()

        for name in method_names:
            metrics_dict = {}
            fpr_tpr = {}
            prob = pred_data_frame[name].values
            fpr, tpr, _ = roc_curve(gt, prob)
            fpr_tpr['fpr'] = fpr
            fpr_tpr['tpr'] = tpr
            # color = color_list[name].values[0]
            auc_value = auc(fpr, tpr)
            metrics_dict['auc'] = auc_value
            metrics_dict['fpr_tpr'] = fpr_tpr
            method_metrics[name] = metrics_dict

            plt.plot(method_metrics[name]['fpr_tpr']['fpr'], method_metrics[name]['fpr_tpr']['tpr'],label="{} (auc={:.3f})".format(name, method_metrics[name]['auc']), c=color_list[0])

        plt.xlabel('False Positive Rate')
        plt.ylabel('True Positive Rate')
        xpoints = ypoints = plt.xlim()
        plt.plot(xpoints, ypoints, color='grey', lw=0.5, scalex=False, scaley=False)
        plt.legend(loc="lower right")
        plt.suptitle('ROC of all trained classifiers')
        plt.savefig(os.path.join(self.save_dir, 'roc_curves_{}.png'.format(data_tag)))

    def plot_confusion_map(self, pred_data_frame, data_tag='train'):
        gt = pred_data_frame[self.label_column_name].values
        prob = pred_data_frame["probs"].values
        # prob = pred_data_frame[self.best_classifier].values
        if self.cutoff is None:
            self.cutoff = self.choose_cutoff(prob, gt)

        pred = (prob >= self.cutoff).astype(np.uint8)
        # if np.sum(pred) != 0:
        cf_matrix = confusion_matrix(gt, pred)
        group_names = ['True Neg', 'False Pos', 'False Neg', 'True Pos']
        group_counts = ["{0: 0.0f}".format(value) for value in cf_matrix.flatten()]
        group_percentages = ["{0:.2%}".format(value) for value in cf_matrix.flatten() / np.sum(cf_matrix)]
        labels = [f"{v1}\n{v2}\n{v3}" for v1, v2, v3 in zip(group_names, group_counts, group_percentages)]
        labels = np.asarray(labels).reshape(2, 2)
        # print('save the confusin matrix with best classifier: {}'.format(self.best_classifier))
        plt.figure()
        sns.heatmap(cf_matrix, annot=labels, fmt='', cmap='Blues')
        plt.savefig(os.path.join(self.save_dir, 'confusion_matrix_for {} set.png'.format(data_tag)))

    def plot_calibration_curve(self, pred_data_frame, data_tag='train'):
        gt = pred_data_frame[self.label_column_name].values
        prob = pred_data_frame[self.best_classifier].values
        prob_true, prob_pred = calibration_curve(gt, prob, n_bins=10)
        plt.figure()
        plt.plot(prob_true, prob_pred, marker='o', c='b', label=self.best_classifier)
        plt.xlabel('Predicted probability')
        plt.ylabel('True probability in each bin')
        xpoints = plt.xlim()
        ypoints = plt.ylim()
        plt.plot(xpoints, ypoints, color='grey', linestyle='--', lw=1, scalex=False,
                 scaley=False)
        plt.legend(loc="upper left")
        plt.suptitle('Calibration plot of {} for {} set'.format(self.best_classifier, data_tag))
        plt.savefig(os.path.join(self.save_dir, 'Calibration plot for {} set.png'.format(data_tag)))

    def report_metrics(self, pred_data_frame):
        """
        :return:
        """
        metrics_list = ['arruracy', 'f1_score', 'Recall', 'Precision', 'AUC', 'Sensitivity', 'Specificity',
                        'Positive prediction', 'Negative prediction', 'Positive llr', 'Negative llr']
        gt = pred_data_frame[self.label_column_name].values
        prob = pred_data_frame[self.best_classifier].values

        # print('current cufoff value is {}'.format(self.cutoff))
        pred = (prob >= self.cutoff).astype(np.uint8)
        accuray = accuracy_score(gt, pred)
        f1 = f1_score(gt, pred)
        recall = recall_score(gt, pred)
        precision = precision_score(gt, pred)
        fpr, tpr, _ = roc_curve(gt, prob)  # here is a list with  all cutoff values
        auc_value = auc(fpr, tpr)
        cf_matrix = confusion_matrix(gt, pred)
        tn, fp, fn, tp = cf_matrix.ravel()
        sensitivity = tp / (tp + fn)
        specificity = tn / (fp + tn)
        ppv = precision  # Positive predictive value
        npv = tn / (fn + tn)  # Negative predictive value
        tpr_ = tp / (tp + fn)  # tpr with current cutoff
        fpr_ = fp / (fp + tn)
        if fpr_ == 0:
            positive_llr = float("NaN")
        else:
            positive_llr = tpr_ / fpr_
        fnr = fn / (tp + fn)
        tnr = tn / (fp + tn)
        if tnr == 0:
            negative_llr = float("NaN")
        else:
            negative_llr = fnr / tnr
        metrics_values = [accuray, f1, recall, precision, auc_value, sensitivity, specificity, ppv, npv, positive_llr,
                          negative_llr]
        return metrics_list, metrics_values

    def plot_decision_curve(self, pred_data_frame, data_tag='train'):
        gt = pred_data_frame[self.label_column_name].values
        prob = pred_data_frame[self.best_classifier].values
        patients_number = len(gt)

        cutoff_list = []
        net_benefit_list = []
        standard_benefit_list = []

        for i in range(0, 100, 1):
            cutoff = i / 100.
            cutoff_list.append(cutoff)
            pred = (prob > cutoff).astype(np.uint8)
            cf_matrix = confusion_matrix(gt, pred)
            tn, fp, fn, tp = cf_matrix.ravel()
            tpr = tp / patients_number
            fpr = fp / patients_number
            net_benefit = tpr - (fpr * (cutoff / (1 - cutoff)))
            net_benefit_list.append(net_benefit)
            standard_benefit = tpr - (1 - fpr) * cutoff / (1 - cutoff)
            standard_benefit_list.append(standard_benefit)
        plt.figure()
        plt.plot(cutoff_list, net_benefit_list, color='red', lw=1, label=self.best_classifier)
        plt.plot(cutoff_list, np.zeros(len(cutoff_list)), color='grey', lw=1, linestyle='--', label='Test None')
        plt.plot(cutoff_list, standard_benefit_list, color='grey', lw=1, label='Test all')
        plt.xlim([0.0, 1.0])
        plt.ylim([-0.1, 1.0])
        plt.xlabel('Threshold Probability')
        plt.ylabel('Net Benefit')
        plt.title('Decision curve for {} set'.format(data_tag))
        plt.legend(loc="upper right")

        plt.savefig(os.path.join(self.save_dir, 'Decision curve for {} set.png'.format(data_tag)))
        # plt.show()

    # def reload_selection_process(self):
    #     feature_selection_process = pd.read_excel(os.path.join(self.save_dir, 'feature_selection_result.xlsx'))
    #     self.document.add_heading('Ⅲ. Feature selection process', level=1)
    #     self.document.add_paragraph('There are {} steps in this process:'.format(len(feature_selection_process)))
    #     for i in range(len(feature_selection_process)):
    #         current_result = feature_selection_process.iloc[i].dropna()
    #         step_name = current_result['Unnamed: 0'].split('_')[-1]
    #         # print(step_name)
    #         num_remained = len(current_result) - 1
    #         self.document.add_paragraph('{}. {}'.format(i + 1, step_name))
    #         self.document.add_paragraph(
    #             'After {} selection, the number of remained features is: {}'.format(step_name, num_remained))
    #         if num_remained < 100:
    #             for n in range(num_remained):
    #                 self.document.add_paragraph('{}'.format(current_result[n]))
    #             self.document.add_paragraph('The heatmap of the model for train set')
    #             self.document.add_picture(
    #                 os.path.join(self.save_dir, 'heatmap_of_train_samples_after_{}.png'.format(step_name)),
    #                 width=Inches(5))
    #             self.document.add_paragraph('The heatmap of the model for test set')
    #             self.document.add_picture(
    #                 os.path.join(self.save_dir, 'heatmap_of_test_samples_after_{}.png'.format(step_name)),
    #                 width=Inches(5))
    #         else:
    #             self.document.add_paragraph('The remained feature is too much to print, please check the excel file!')
    #         if step_name == 'Lasso':
    #             self.document.add_paragraph('lasso loss plot for train set')
    #             self.document.add_picture(os.path.join(self.save_dir, '{}_loss_plot.png'.format(step_name)),
    #                                       width=Inches(5))
    #             self.document.add_paragraph('lasso coefficients plot for train set')
    #             self.document.add_picture(os.path.join(self.save_dir, '{}_coefficient_plot.png'.format(step_name)),
    #                                       width=Inches(5))
    #             self.document.add_paragraph('feature importance after lasso is:')
    #             self.document.add_picture(os.path.join(self.save_dir, 'lasso_feature_importance.png'), width=Inches(5))
    #         if step_name == 'RandomForestImportance':
    #             self.document.add_paragraph('feature importance from random forest')
    #             self.document.add_picture(os.path.join(self.save_dir, 'feature_importance_with_random_forest.png'),
    #                                       width=Inches(5))


    def run_analysis(self):
        # select best classifier

        assert self.test_pred_data_frame is not None, "test prediction data must not be None "

        if self.do_feature_selection == True:
            # self.reload_selection_process()
            pass
        else:
            self.document.add_heading('Ⅲ. Feature selection process', level=1)
            self.document.add_paragraph("You don't do feature selection.")

        self.plot_roc_curves(self.test_pred_data_frame,data_tag="test")
        if self.train_pred_data_frame is not None:
            self.plot_roc_curves(self.train_pred_data_frame, data_tag="train")

        # plot consufion matrix
        # self.document.add_heading('Ⅳ. analysis part', level=1)

        self.plot_confusion_map(self.test_pred_data_frame, data_tag='test')
        # self.document.add_paragraph('The best cutoff with maxmizing Youden Index is {}'.format(self.cutoff))
        # self.document.add_paragraph('1. confusion matrixes')
        # self.document.add_paragraph('confusion matrix for test set')
        # self.document.add_picture(os.path.join(self.save_dir, 'confusion_matrix_for test set.png'),
        #                           width=Inches(5))
        if self.train_pred_data_frame is not None:
            self.plot_confusion_map(self.train_pred_data_frame, data_tag='train')
        #     self.document.add_paragraph('confusion matrix for train set')
        #     self.document.add_picture(os.path.join(self.save_dir, 'confusion_matrix_for train set.png'),
        #                               width=Inches(5))
        #
        # metrics_list, metrics_values = self.report_metrics(self.test_pred_data_frame)
        # metrics_values = np.array(metrics_values)
        # metrics_values_all = metrics_values[np.newaxis]
        #
        # if self.train_pred_data_frame is not None:
        #     metrics_list_train, metrics_values_train = self.report_metrics(self.train_pred_data_frame)
        #     metrics_values_train = np.array(metrics_values_train)
        #     metrics_values_all = np.concatenate([metrics_values_all, metrics_values_train[np.newaxis]], axis=0)
        # if metrics_values_all.shape[0] == 2:
        #     columns = ['test', 'train']
        # else:
        #     columns = ['test']
        #
        # metrics_df = pd.DataFrame(metrics_values_all.transpose(), index=metrics_list, columns=columns)
        # metrics_df.to_excel(os.path.join(self.save_dir, 'metrics_report.xlsx'))
        #
        # self.document.add_paragraph('2. classification metrics')
        # print_df = pd.read_excel(os.path.join(self.save_dir, 'metrics_report.xlsx'))
        # table = self.document.add_table(print_df.shape[0] + 1, print_df.shape[1], style='Table Grid')
        # for j in range(print_df.shape[1]):
        #     table.cell(0, j).text = print_df.columns[j]
        #
        # for i in range(print_df.shape[0]):
        #     for j in range(print_df.shape[1]):
        #         table.cell(i + 1, j).text = str(print_df.values[i, j])
        #
        # #
        # self.document.add_paragraph('3. calibration curves')
        # self.plot_calibration_curve(self.test_pred_data_frame, data_tag='test')
        # self.document.add_paragraph('calibration curve for test set')
        # self.document.add_picture(os.path.join(self.save_dir, 'Calibration plot for test set.png'), width=Inches(5))
        # if self.train_pred_data_frame is not None:
        #     self.plot_calibration_curve(self.train_pred_data_frame, data_tag='train')
        #     self.document.add_paragraph('calibration curve for train set')
        #     self.document.add_picture(os.path.join(self.save_dir, 'Calibration plot for train set.png'),
        #                               width=Inches(5))
        #
        # #
        # self.document.add_paragraph('4. decision curve analysis')
        # self.plot_decision_curve(self.test_pred_data_frame, data_tag='test')
        # self.document.add_paragraph('decision curve for test set')
        # self.document.add_picture(os.path.join(self.save_dir, 'Decision curve for test set.png'), width=Inches(5))
        #
        # if self.train_pred_data_frame is not None:
        #     self.plot_decision_curve(self.train_pred_data_frame, data_tag='train')
        #     self.document.add_paragraph('decision curve for train set')
        #     self.document.add_picture(os.path.join(self.save_dir, 'Decision curve for train set.png'), width=Inches(5))

        # self.document.save(os.path.join(self.save_dir, 'result_report_{}.docx'.format(str(time()).split('.')[0])))

